"""文件输入解析。"""

import base64
import csv
import json
import mimetypes
import re
from dataclasses import dataclass
from pathlib import Path

from langchain_core.messages import HumanMessage

from agent_app.config import MAX_FILE_SIZE_MB


TEXT_EXTENSIONS = {".txt", ".md", ".json", ".csv"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}
DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".xlsx"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | IMAGE_EXTENSIONS | DOCUMENT_EXTENSIONS
FILE_REF_PATTERN = re.compile(r"@(?P<path>(?:\"[^\"]+\"|'[^']+'|\S+))")


@dataclass
class FileParseResult:
    """单个文件解析结果。"""

    path: str
    kind: str
    content: str = ""
    data_url: str = ""
    error: str = ""


def parse_user_input(user_input: str) -> tuple[str, list[FileParseResult]]:
    """解析用户输入中的 @文件路径。"""
    file_results = []

    def replace_match(match):
        raw_path = match.group("path").strip("\"'")
        result = parse_file(raw_path)
        file_results.append(result)
        return f"[文件: {raw_path}]"

    text = FILE_REF_PATTERN.sub(replace_match, user_input)
    return text, file_results


def parse_file(path: str) -> FileParseResult:
    """按扩展名解析文件。"""
    file_path = Path(path).expanduser()
    if not file_path.is_file():
        return FileParseResult(path=str(file_path), kind="error", error="文件不存在或不是普通文件。")

    max_file_size_bytes = int(MAX_FILE_SIZE_MB * 1024 * 1024)
    try:
        file_size = file_path.stat().st_size
    except OSError as exc:
        return FileParseResult(path=str(file_path), kind="error", error=f"无法读取文件信息：{exc}")
    if max_file_size_bytes >= 0 and file_size > max_file_size_bytes:
        return FileParseResult(
            path=str(file_path),
            kind="error",
            error=f"文件过大：{_format_size(file_size)}，超过限制 {_format_size(max_file_size_bytes)}。",
        )

    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        return FileParseResult(path=str(file_path), kind="error", error=f"不支持的文件类型：{suffix}")

    try:
        if suffix in {".txt", ".md"}:
            return FileParseResult(path=str(file_path), kind="text", content=file_path.read_text(encoding="utf-8"))
        if suffix == ".json":
            return FileParseResult(path=str(file_path), kind="text", content=_parse_json(file_path))
        if suffix == ".csv":
            return FileParseResult(path=str(file_path), kind="text", content=_parse_csv(file_path))
        if suffix == ".pdf":
            return FileParseResult(path=str(file_path), kind="document", content=_parse_pdf(file_path))
        if suffix == ".docx":
            return FileParseResult(path=str(file_path), kind="document", content=_parse_docx(file_path))
        if suffix == ".xlsx":
            return FileParseResult(path=str(file_path), kind="document", content=_parse_xlsx(file_path))
        if suffix in IMAGE_EXTENSIONS:
            return FileParseResult(path=str(file_path), kind="image", data_url=_image_to_data_url(file_path))
    except ImportError as exc:
        return FileParseResult(path=str(file_path), kind="error", error=f"缺少解析依赖：{exc.name}")
    except UnicodeDecodeError:
        return FileParseResult(path=str(file_path), kind="error", error="文件不是 UTF-8 文本，暂无法解析。")
    except Exception as exc:
        return FileParseResult(path=str(file_path), kind="error", error=f"文件解析失败：{exc}")

    return FileParseResult(path=str(file_path), kind="error", error="未知文件解析错误。")


def build_human_message(text: str, file_results: list[FileParseResult]) -> HumanMessage:
    """根据文本和文件解析结果构造 HumanMessage。"""
    image_results = [result for result in file_results if result.kind == "image" and result.data_url]
    text_content = _build_text_context(text, file_results)

    if not image_results:
        return HumanMessage(content=text_content)

    content = [{"type": "text", "text": text_content}]
    for result in image_results:
        content.append({"type": "image_url", "image_url": {"url": result.data_url}})
    return HumanMessage(content=content)


def _build_text_context(text: str, file_results: list[FileParseResult]) -> str:
    """把文件解析结果拼接到用户文本后。"""
    parts = [text]
    for result in file_results:
        if result.error:
            parts.append(f"\n\n[文件解析失败]\n路径：{result.path}\n原因：{result.error}")
        elif result.kind == "image":
            parts.append(f"\n\n[图片文件]\n路径：{result.path}\n说明：图片已作为多模态输入附加。")
        else:
            parts.append(f"\n\n[文件内容]\n路径：{result.path}\n类型：{result.kind}\n内容：\n{result.content}")
    return "\n".join(parts).strip()


def _parse_json(file_path: Path) -> str:
    data = json.loads(file_path.read_text(encoding="utf-8"))
    return json.dumps(data, ensure_ascii=False, indent=2)


def _parse_csv(file_path: Path) -> str:
    with file_path.open("r", encoding="utf-8", newline="") as file:
        rows = list(csv.reader(file))
    return "\n".join([" | ".join(row) for row in rows])


def _parse_pdf(file_path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"## 第 {index} 页\n{text.strip()}")

    if not pages:
        return "PDF 未提取到文本内容。可能是扫描版 PDF，请转成图片后再使用图片输入。"
    return "\n\n".join(pages)


def _parse_docx(file_path: Path) -> str:
    from docx import Document

    document = Document(str(file_path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"## 表格 {table_index}")
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))

    return "\n".join(parts).strip() or "DOCX 未提取到文本内容。"


def _parse_xlsx(file_path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(str(file_path), read_only=True, data_only=True)
    parts = []
    for sheet in workbook.worksheets:
        parts.append(f"## 工作表：{sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(values):
                parts.append(" | ".join(values))

    return "\n".join(parts).strip() or "XLSX 未提取到表格内容。"


def _image_to_data_url(file_path: Path) -> str:
    mime_type = mimetypes.guess_type(str(file_path))[0] or "application/octet-stream"
    encoded = base64.b64encode(file_path.read_bytes()).decode("ascii")
    return f"data:{mime_type};base64,{encoded}"


def _format_size(size_bytes: int) -> str:
    """格式化文件体积。"""
    size_mb = size_bytes / 1024 / 1024
    if size_mb >= 1:
        return f"{size_mb:.2f} MB"
    return f"{size_bytes / 1024:.2f} KB"
